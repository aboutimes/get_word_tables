import os,sys
import docx
import xlwt
import datetime

from docx import Document #导入库

rootdir = os.getcwd() + "/doc"  #文件路径
file_type1 ='doc'#指定文件类型
file_type2 ='docx'#指定文件类型
list = os.listdir(rootdir) #列出文件夹下所有的目录与文件
for f in list:
  file_suffix = f.split('.')
  file_suffix = file_suffix[-1] # 取后缀
  if (file_type1 != file_suffix and file_type2 != file_suffix):
    list.remove(f)  # 删除非指定文件类型
# 输出    
new_excel = xlwt.Workbook() #创建工作簿
sheet1 = new_excel.add_sheet(u'sheet1',cell_overwrite_ok=True) #创建sheet
for i in range(0,len(list)):
  file_dir = os.path.join(rootdir,list[i])
  file_suffix = os.path.splitext(file_dir)[1]
  if (os.path.isfile(file_dir)):  # 如果是文件夹就继续打开
    document = Document(file_dir) #读入文件
    tables = document.tables #获取文件中的表格集
    table = tables[0] #获取文件中的第一个表格
    # 1.表头
    if i==0:
      name_title = table.cell(0,0).text #cell( )  这是表格里面单元格的位置
      sex_title = table.cell(0,2).text
      birth_title = table.cell(0,4).text
      data_title = [name_title,sex_title,birth_title]
      # 输出表头
      for j in range(len(data_title)):
        sheet1.write(i,j,data_title[j])  #第一个是写入哪一行,第二个写入参数的列, 第三个是要写入的数据
    # 2.表体
    name = table.cell(0,1).text #cell( )  这是表格里面单元格的位置
    sex = table.cell(0,3).text
    birth = table.cell(0,5).text
    data=[] #创建一个空列表，用append将表格里取出的数据加入到这个列表
    # data.append(Document(path).paragraphs[1].text)  #这里是读取了word中的第一段的内容
    data.append(name)
    data.append(sex)
    data.append(birth)
    # 输出表体
    data_length=range(len(data))
    for k in data_length:
      sheet1.write(i+1,k,data[k])  #第一个是写入哪一行,第二个写入参数的列, 第三个是要写入的数据
file_name = datetime.datetime.now().strftime('%Y-%m-%d %H.%M.%S') + ".xls"  # 定义保存的文件名
new_excel.save(file_name)  #保存文件
